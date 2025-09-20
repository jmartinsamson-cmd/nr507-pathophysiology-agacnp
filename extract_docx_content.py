import os
import json
from pathlib import Path
import re

# Install python-docx if needed
try:
    from docx import Document
    print("python-docx is available")
except ImportError:
    print("Installing python-docx...")
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    print("python-docx installed and imported")

def extract_questions_from_docx(docx_path):
    """Extract questions from a docx file"""
    try:
        doc = Document(docx_path)

        # Get all text content from the document
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())

        # Also check tables for content
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())

        # Join all text
        content = '\n'.join(full_text)

        # Extract questions using patterns
        questions = []

        # Pattern to match numbered questions
        question_pattern = r'(\d+)\.\s*(.+?)(?=\n\d+\.|$)'
        matches = re.findall(question_pattern, content, re.DOTALL)

        for match in matches:
            question_num = match[0]
            question_text = match[1].strip()

            # Split question from options/answer
            lines = question_text.split('\n')
            question_main = lines[0].strip()

            # Look for options (A, B, C, D or 1, 2, 3, 4)
            options = []
            correct_answers = []
            rationale = ""

            for i, line in enumerate(lines[1:], 1):
                line = line.strip()
                if re.match(r'^[A-D]\)', line) or re.match(r'^[A-D]\.', line):
                    options.append(line[2:].strip())
                elif re.match(r'^[1-4]\)', line) or re.match(r'^[1-4]\.', line):
                    options.append(line[2:].strip())
                elif 'correct answer' in line.lower() or 'answer:' in line.lower():
                    # Extract correct answer
                    answer_match = re.search(r'[A-D]|[1-4]', line)
                    if answer_match:
                        ans = answer_match.group()
                        if ans.isdigit():
                            correct_answers.append(int(ans) - 1)
                        else:
                            correct_answers.append(ord(ans.upper()) - ord('A'))
                elif 'rationale' in line.lower() or 'explanation' in line.lower():
                    rationale = line

            if question_main and options:
                questions.append({
                    'question_number': int(question_num),
                    'question_text': question_main,
                    'options': options,
                    'correct_answers': correct_answers if correct_answers else [0],
                    'rationale': rationale
                })

        return questions, content

    except Exception as e:
        print(f"Error reading {docx_path}: {str(e)}")
        return [], ""

# Define file paths
base_path = r"C:\Users\jsamb\OneDrive\Desktop\Exampee"
chapters = [28, 29, 30, 32, 34, 35, 37, 38]

# Extract content from all docx files
all_extracted_content = {}

for chapter in chapters:
    docx_path = f"{base_path}\\Chapter {chapter}.docx"
    print(f"\nProcessing Chapter {chapter}...")

    if os.path.exists(docx_path):
        questions, raw_content = extract_questions_from_docx(docx_path)
        all_extracted_content[chapter] = {
            'questions': questions,
            'raw_content': raw_content[:2000],  # First 2000 chars for inspection
            'question_count': len(questions)
        }
        print(f"  Extracted {len(questions)} questions")
    else:
        print(f"  File not found: {docx_path}")
        all_extracted_content[chapter] = {
            'questions': [],
            'raw_content': "",
            'question_count': 0
        }

# Save extracted content for inspection
with open(f"{base_path}\\extracted_docx_content.json", 'w', encoding='utf-8') as f:
    json.dump(all_extracted_content, f, indent=2, ensure_ascii=False)

print(f"\nExtraction complete. Results saved to extracted_docx_content.json")
print("\nSummary:")
for chapter in chapters:
    count = all_extracted_content[chapter]['question_count']
    print(f"  Chapter {chapter}: {count} questions")